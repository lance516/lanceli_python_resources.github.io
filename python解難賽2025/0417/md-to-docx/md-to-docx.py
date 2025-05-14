# this program reads all the .md files in the md folder and converts them to .docx files in the docx folder
# It reads each .md line by line, convert headings to headings
# and normal lines to paragraphs.
# use os, glob, docx.
import os
import glob
from docx import Document
import markdown
from markdown.extensions.tables import TableExtension

def convert_md_to_docx(input_file, output_dir='docx'):
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate output filename
    base_name = os.path.basename(input_file)
    output_file = os.path.join(output_dir, os.path.splitext(base_name)[0] + '.docx')
    
    # Read markdown content
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML with table support
    html_content = markdown.markdown(md_content, extensions=['tables'])
    
    # Create a new Word document
    doc = Document()
    
    # Add the HTML content with headings and paragraphs
    for block in html_content.splitlines():
        if block.startswith('# '):
            doc.add_heading(block[2:], 1)
        elif block.startswith('## '):
            doc.add_heading(block[3:], 2)
        elif block.startswith('### '):
            doc.add_heading(block[4:], 3)
        elif block.startswith('#### '):
            doc.add_heading(block[5:], 4)
        elif block.startswith('##### '):
            doc.add_heading(block[6:], 5)
        elif block.startswith('###### '):
            doc.add_heading(block[7:], 6)
        else:
            doc.add_paragraph(block)
    
    # Save the document
    doc.save(output_file)
    print(f'Successfully converted {input_file} to {output_file}')

def main():
    # Process all .md files in the md directory
    md_files = glob.glob('md/*.md')
    
    if not md_files:
        print('No markdown files found in the md directory')
        return
    
    for md_file in md_files:
        convert_md_to_docx(md_file)

if __name__ == '__main__':
    main()
