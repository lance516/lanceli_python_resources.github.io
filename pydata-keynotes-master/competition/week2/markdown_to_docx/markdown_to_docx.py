"""
# 基礎 Markdown 轉 DOCX

- 檔取文字當檔做基本分析
- 將分析結果轉化成 docx 檔案
"""

import os
from docx import Document


def read_markdown(filename):

    data = []
    new_paragraph = False

    with open(filename, 'r') as file:
        for line in file:

            last_index = len(data)-1
            last_line = data[last_index] if last_index >= 0 else None
            line = line.strip("\n")

            if line.startswith("#"):
                # check if line starts with header indicator (#)
                data.append(('heading', line.replace('#', '').strip()))
            elif line == "---" or line == "----":
                # check if line starts with seperator (--- or ----)
                data.append(('line_break', ''))
            else:
                # paragraph
                if line == "":
                    # check if line empty, indicate next paragraph as new paragraph
                    new_paragraph = True
                else:
                    if new_paragraph or last_line[0] != "paragraph":
                        # if indicated new_paragraph or previous item is not paragraph, put a new paragraph
                        new_paragraph = False
                        data.append(('paragraph', line))
                    else:
                        # else merge with previous paragraph
                        data.pop()
                        data.append(('paragraph', last_line[1] + line))

    return data


def write_docx(filename, parsed_data):
    doc = Document()

    for line in parsed_data:
        if line[0] == 'heading':
            doc.add_heading(line[1], level=1)
        elif line[0] == 'line_break':
            doc.add_page_break()
        elif line[0] == "paragraph":
            doc.add_paragraph(line[1])

    doc.save(filename)


def main():
    path = os.path.dirname(__file__)
    input_dir = os.path.join(path, "inputs")
    output_dir = os.path.join(path, "outputs")
    # filename = os.path.join(path, "inputs", "sample.md")
    # out_filename = os.path.join(path, "outputs", "sample.docx")

    for filename in os.listdir(input_dir):
        # get the full path
        file_path = os.path.join(input_dir, filename)
        # get the filename from sample.md (sample)
        out_filename, out_fileext = os.path.splitext(filename)
        out_filepath = os.path.join(output_dir, out_filename + ".docx")

        parsed = read_markdown(file_path)
        write_docx(out_filepath, parsed)
        # print(parsed)


if __name__ == "__main__":
    main()
