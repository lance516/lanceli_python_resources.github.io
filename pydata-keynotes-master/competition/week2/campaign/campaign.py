"""
# 題目: 圖像檔案分類成 Word 文檔

- 使用 os.listdir 列出檔案
- 使用 re.match 解拆檔案名
- 使用 docx 加入圖片
"""

import os
import re
from docx import Document


def find_photos(path):

    groups = {}

    for filename in sorted(os.listdir(path)):
        name_only = os.path.splitext(filename)
        matches = re.match("^(.*)-(\d*)$", name_only[0])
        # print(matches)

        if matches != None:
            prefix = matches.group(1)  # first match group (.*)

            if groups.get(prefix) != None:
                groups[prefix].append(filename)
            else:
                groups[prefix] = [filename]

    return groups


def write_docx(input_path, output_path, groups):

    for key, items in groups.items():

        doc = Document()

        for item in items:
            doc.add_picture(os.path.join(input_path, item))
            doc.add_paragraph('')

        doc.save(os.path.join(output_path, f"{key}.docx"))


def main():
    path = os.path.dirname(__file__)
    input_path = os.path.join(path, "inputs")
    output_path = os.path.join(path, "outputs")

    groups = find_photos(input_path)
    write_docx(input_path, output_path, groups)


if __name__ == "__main__":
    main()
